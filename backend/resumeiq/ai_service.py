import re
from PyPDF2 import PdfReader
from docx import Document
import json


DOMAIN_KEYWORDS = {
    'software_engineering': {
        'keywords': ['Python', 'JavaScript', 'React', 'Node.js', 'Docker', 'AWS', 'Git', 'SQL', 'API', 'CI/CD', 'Agile', 'REST', 'GraphQL', 'TypeScript', 'MongoDB', 'PostgreSQL'],
        'roles': ['Software Engineer', 'Full Stack Developer', 'Backend Developer', 'Frontend Developer']
    },
    'accounting_finance': {
        'keywords': ['Accounting', 'Financial Reporting', 'GST', 'Tally Prime', 'Reconciliation', 'Audit', 'Budgeting', 'Forecasting', 'Excel', 'QuickBooks', 'SAP', 'Taxation', 'Financial Analysis'],
        'roles': ['Accountant', 'Financial Analyst', 'Auditor', 'Finance Manager']
    },
    'marketing': {
        'keywords': ['SEO', 'Google Analytics', 'Content Marketing', 'Social Media', 'Email Marketing', 'CRM', 'Digital Marketing', 'Brand Strategy', 'Market Research', 'PPC', 'Conversion Optimization'],
        'roles': ['Marketing Manager', 'Digital Marketer', 'SEO Specialist', 'Content Strategist']
    },
    'healthcare': {
        'keywords': ['Patient Care', 'EHR', 'HIPAA', 'Clinical Research', 'Medical Terminology', 'EMR', 'Healthcare Management', 'Nursing', 'Patient Safety', 'Regulatory Compliance'],
        'roles': ['Registered Nurse', 'Healthcare Administrator', 'Medical Coder', 'Clinical Coordinator']
    },
    'education': {
        'keywords': ['Curriculum Development', 'Instructional Design', 'Classroom Management', 'Learning Management Systems', 'Assessment', 'Differentiated Instruction', 'Educational Technology', 'Student Engagement'],
        'roles': ['Teacher', 'Instructional Designer', 'Education Coordinator', 'Curriculum Specialist']
    }
}


def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = '\n'.join([page.extract_text() or '' for page in reader.pages])
    return text


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return text


def detect_domain(resume_text, job_desc=''):
    combined = (resume_text + ' ' + job_desc).lower()
    scores = {}
    for domain, data in DOMAIN_KEYWORDS.items():
        count = sum(1 for kw in data['keywords'] if kw.lower() in combined)
        scores[domain] = count
    best_domain = max(scores, key=scores.get) if scores else 'software_engineering'
    return best_domain if scores[best_domain] > 0 else 'software_engineering'


def analyze_resume(resume_text, job_desc=''):
    domain = detect_domain(resume_text, job_desc)
    domain_kws = DOMAIN_KEYWORDS[domain]['keywords']
    lower_text = resume_text.lower()

    matched = [kw for kw in domain_kws if kw.lower() in lower_text]
    missing = [kw for kw in domain_kws if kw.lower() not in lower_text]

    # Generate detailed missing keyword objects
    missing_detailed = []
    for kw in missing:
        missing_detailed.append({
            'keyword': kw,
            'why_it_matters': f'{kw} is highly relevant for roles in {domain.replace("_", " ").title()}. It demonstrates expertise in key areas employers look for.',
            'example_bullet': f'Implemented {kw} to streamline processes and improve efficiency across the team'
        })

    # Calculate scores
    keyword_score = min(100, int((len(matched) / len(domain_kws)) * 40)) if domain_kws else 20
    len_score = min(30, int(len(resume_text) / 50))
    structure_score = 20
    if 'experience' in lower_text or 'work history' in lower_text:
        structure_score += 5
    if 'education' in lower_text:
        structure_score += 5

    total = keyword_score + len_score + structure_score
    ats_score = min(100, total)

    score_breakdown = {
        'Keyword Relevance': keyword_score,
        'Content Depth': len_score,
        'Structure & Formatting': structure_score,
    }

    # Strengths & Weaknesses
    strengths = []
    if len(matched) > 3:
        strengths.append(f"Strong keyword coverage with {len(matched)} relevant terms matched")
    if len(resume_text) > 1500:
        strengths.append("Good content depth with detailed descriptions")
    if 'experience' in lower_text and 'education' in lower_text:
        strengths.append("Well-structured with clear Experience and Education sections")

    weaknesses = []
    if len(missing) > 3:
        weaknesses.append(f"Missing {len(missing)} key domain-specific keywords")
    if len(resume_text) < 800:
        weaknesses.append("Resume content is too brief; add more details about your experience")
    if 'achievements' not in lower_text and 'results' not in lower_text:
        weaknesses.append("Lack of quantifiable achievements or results")

    # AI Suggestions (only optional tech clearly marked)
    ai_suggestions = [
        "Use action verbs like 'Led', 'Developed', 'Implemented' to start bullet points",
        "Quantify achievements with numbers and metrics where possible",
        "Keep formatting consistent and clean for optimal ATS readability"
    ]

    # Optional learning paths (clearly marked)
    optional_paths = [f"Optional Learning Path: Consider exploring {kw} to expand your skill set" for kw in missing[:2]]
    ai_suggestions.extend(optional_paths)

    # Generate improved resume
    improved = generate_improved_resume(resume_text, matched, domain)

    return {
        'ats_score': ats_score,
        'score_breakdown': score_breakdown,
        'matched_kw': matched,
        'missing_kw': missing_detailed,
        'ai_suggestions': ai_suggestions,
        'strengths': strengths,
        'weaknesses': weaknesses,
        'improved_resume': improved,
        'domain': domain
    }


def generate_improved_resume(original_text, matched_kws, domain):
    lines = original_text.split('\n')
    cleaned = [l.strip() for l in lines if l.strip()]
    
    name = cleaned[0] if cleaned else "Your Name"
    contact_info = ""
    
    possible_contact_lines = []
    for line in cleaned[1:5]:
        if '@' in line or 'linkedin' in line.lower() or 'github' in line.lower() or any(d in line for d in ['(', '-', ')']):
            possible_contact_lines.append(line)
    
    if possible_contact_lines:
        contact_info = " | ".join(possible_contact_lines)
    
    sections = []
    current_title = None
    current_content = []
    for line in cleaned[1:]:
        if contact_info and line in possible_contact_lines:
            continue
            
        lower_line = line.lower()
        section_keywords = [
            'experience', 'education', 'skills', 'summary', 'projects', 
            'objective', 'work history', 'work experience', 'academic background',
            'certifications', 'awards', 'achievements', 'technical skills'
        ]
        if any(kw in lower_line for kw in section_keywords):
            if current_title and current_content:
                sections.append((current_title, current_content))
            current_title = line.upper() if line else None
            current_content = []
        else:
            current_content.append(line)
    if current_title and current_content:
        sections.append((current_title, current_content))
    
    improved_parts = []
    improved_parts.append(name.upper())
    if contact_info:
        improved_parts.append(contact_info)
    improved_parts.append("")
    
    professional_summary = None
    experience = None
    technical_skills = None
    projects = None
    certifications = None
    education = None
    achievements = None
    other_sections = []
    
    for title, content in sections:
        t_low = title.lower()
        if 'summary' in t_low or 'objective' in t_low:
            if not professional_summary:
                professional_summary = (title, content)
        elif 'experience' in t_low or 'work' in t_low:
            if not experience:
                experience = (title, content)
        elif 'skills' in t_low or 'technical' in t_low:
            if not technical_skills:
                technical_skills = (title, content)
        elif 'projects' in t_low:
            if not projects:
                projects = (title, content)
        elif 'certifications' in t_low:
            if not certifications:
                certifications = (title, content)
        elif 'education' in t_low or 'academic' in t_low:
            if not education:
                education = (title, content)
        elif 'achievements' in t_low or 'awards' in t_low:
            if not achievements:
                achievements = (title, content)
        else:
            other_sections.append((title, content))
    
    # HEADER done
    
    # PROFESSIONAL SUMMARY
    if professional_summary:
        improved_parts.append("PROFESSIONAL SUMMARY")
        improved_parts.append("-" * 40)
        for line in professional_summary[1]:
            if line.strip():
                improved_parts.append(line.strip())
        improved_parts.append("")
    
    # EXPERIENCE
    if experience:
        improved_parts.append("EXPERIENCE")
        improved_parts.append("-" * 40)
        for line in experience[1]:
            if line.strip():
                improved_parts.append(line.strip())
        improved_parts.append("")
    
    # TECHNICAL SKILLS
    improved_parts.append("TECHNICAL SKILLS")
    improved_parts.append("-" * 40)
    if technical_skills:
        for line in technical_skills[1]:
            if line.strip():
                improved_parts.append(line.strip())
    else:
        categories = {
            "Languages": [],
            "Frameworks": [],
            "Databases": [],
            "Tools": [],
            "Concepts": [],
            "Cloud": []
        }
        tech_keywords = {
            "software_engineering": {
                "Languages": ["Python", "JavaScript", "TypeScript", "Java", "C++", "Go", "Rust", "Ruby", "PHP"],
                "Frameworks": ["React", "Node.js", "Django", "Spring Boot", "Vue", "Angular", "Express", "Flask"],
                "Databases": ["MySQL", "PostgreSQL", "MongoDB", "SQL", "Redis", "Oracle"],
                "Tools": ["Git", "Docker", "AWS", "CI/CD", "Jira", "VS Code", "Postman"],
                "Concepts": ["REST API", "GraphQL", "Agile", "DevOps", "OOP", "MVC"],
                "Cloud": ["AWS", "Azure", "GCP", "S3", "EC2"]
            },
            "accounting_finance": {
                "Languages": ["Excel", "SQL"],
                "Frameworks": [],
                "Databases": ["QuickBooks", "SAP"],
                "Tools": ["Tally Prime", "Excel", "Power BI", "Tableau"],
                "Concepts": ["Financial Reporting", "Reconciliation", "Audit", "Budgeting", "Taxation"],
                "Cloud": []
            }
        }
        for kw in matched_kws:
            categorized = False
            domain_cats = tech_keywords.get(domain, tech_keywords["software_engineering"])
            for cat, cat_kw in domain_cats.items():
                if kw in cat_kw or any(c in kw.lower() for c in cat.lower().split()):
                    categories[cat].append(kw)
                    categorized = True
                    break
            if not categorized:
                categories["Tools"].append(kw)
        
        # Format with proper alignment
        max_cat_length = max(len(cat) for cat in categories if categories[cat]) if any(categories.values()) else 10
        for cat, skills in categories.items():
            if skills:
                padding = " " * (max_cat_length - len(cat))
                improved_parts.append(f"{cat}{padding} : {', '.join(skills)}")
    improved_parts.append("")
    
    # PROJECTS
    if projects:
        improved_parts.append("PROJECTS")
        improved_parts.append("-" * 40)
        for line in projects[1]:
            if line.strip():
                improved_parts.append(line.strip())
        improved_parts.append("")
    
    # CERTIFICATIONS
    if certifications:
        improved_parts.append("CERTIFICATIONS")
        improved_parts.append("-" * 40)
        for line in certifications[1]:
            if line.strip():
                improved_parts.append(line.strip())
        improved_parts.append("")
    
    # EDUCATION
    if education:
        improved_parts.append("EDUCATION")
        improved_parts.append("-" * 40)
        for line in education[1]:
            if line.strip():
                improved_parts.append(line.strip())
        improved_parts.append("")
    
    # ACHIEVEMENTS
    if achievements:
        improved_parts.append("ACHIEVEMENTS")
        improved_parts.append("-" * 40)
        for line in achievements[1]:
            if line.strip():
                improved_parts.append(line.strip())
        improved_parts.append("")
    
    # OTHER SECTIONS
    for title, content in other_sections:
        improved_parts.append(title)
        improved_parts.append("-" * 40)
        for line in content:
            if line.strip():
                improved_parts.append(line.strip())
        improved_parts.append("")
    
    return '\n'.join(improved_parts).strip()


def generate_pdf_content(resume_content):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.lib.colors import gray, black
    from io import BytesIO

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    
    name_style = ParagraphStyle(
        'NameStyle',
        parent=styles['Heading1'],
        fontSize=26,
        leading=30,
        alignment=TA_CENTER,
        spaceAfter=8,
        fontName='Helvetica-Bold',
        textColor=black
    )
    
    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['BodyText'],
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    section_style = ParagraphStyle(
        'SectionStyle',
        parent=styles['Heading2'],
        fontSize=15,
        leading=18,
        spaceAfter=4,
        spaceBefore=20,
        fontName='Helvetica-Bold'
    )
    
    project_title_style = ParagraphStyle(
        'ProjectTitleStyle',
        parent=styles['Heading3'],
        fontSize=13,
        leading=16,
        spaceAfter=4,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['BodyText'],
        fontSize=10.5,
        leading=13,
        spaceAfter=4
    )

    story = []
    lines = resume_content.split('\n')
    
    in_projects_section = False
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        if stripped.startswith('---') or stripped.startswith('===') or all(c == '-' for c in stripped):
            story.append(HRFlowable(width="100%", thickness=1, color=gray, spaceBefore=0, spaceAfter=10))
        elif stripped:
            if i == 0:
                story.append(Paragraph(f'<b>{stripped}</b>', name_style))
            elif (i == 1 or i == 2) and '@' in stripped:
                story.append(Paragraph(stripped, contact_style))
            else:
                is_section_heading = (
                    stripped.isupper() and 
                    len(stripped) > 3 and 
                    not any(char.isdigit() for char in stripped) and
                    not any(c in stripped for c in ['@', '.com', 'http', '|'])
                )
                
                common_sections = [
                    'PROFESSIONAL SUMMARY', 'EXPERIENCE', 'EDUCATION', 'TECHNICAL SKILLS', 
                    'PROJECTS', 'CERTIFICATIONS', 'AWARDS', 'ACHIEVEMENTS', 'WORK HISTORY',
                    'WORK EXPERIENCE', 'ACADEMIC BACKGROUND', 'OBJECTIVE', 'KEY SKILLS'
                ]
                is_section_heading = is_section_heading or stripped.upper() in common_sections
                
                if is_section_heading:
                    story.append(Paragraph(f'<b>{stripped}</b>', section_style))
                    if stripped.upper() == "PROJECTS":
                        in_projects_section = True
                    else:
                        in_projects_section = False
                else:
                    # Check if it's a project title (not a bullet, not a section, in projects section)
                    is_project_title = (
                        in_projects_section and 
                        not stripped.startswith(('•', '-', '*', '○')) and 
                        not stripped.startswith(('http', 'www', 'Tech Stack', 'Languages', 'Frameworks', 'Databases', 'Tools', 'Concepts', 'Cloud'))
                    )
                    
                    if is_project_title:
                        story.append(Paragraph(f'<b>{stripped}</b>', project_title_style))
                    else:
                        story.append(Paragraph(stripped, body_style))
        else:
            story.append(Spacer(1, 0.05*inch))

    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_docx_content(resume_content):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO

    def add_horizontal_line(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C0C0C0')
        pBdr.append(bottom)
        pPr.append(pBdr)

    buffer = BytesIO()
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = resume_content.split('\n')
    in_projects_section = False
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        if stripped.startswith('---') or stripped.startswith('===') or all(c == '-' for c in stripped):
            p = doc.add_paragraph()
            add_horizontal_line(p)
            p.paragraph_format.space_after = Pt(10)
        elif stripped:
            p = doc.add_paragraph()
            
            if i == 0:
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(26)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.space_after = Pt(8)
            elif (i == 1 or i == 2) and '@' in stripped:
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.space_after = Pt(20)
            else:
                is_section_heading = (
                    stripped.isupper() and 
                    len(stripped) > 3 and 
                    not any(char.isdigit() for char in stripped) and
                    not any(c in stripped for c in ['@', '.com', 'http', '|'])
                )
                
                common_sections = [
                    'PROFESSIONAL SUMMARY', 'EXPERIENCE', 'EDUCATION', 'TECHNICAL SKILLS', 
                    'PROJECTS', 'CERTIFICATIONS', 'AWARDS', 'ACHIEVEMENTS', 'WORK HISTORY',
                    'WORK EXPERIENCE', 'ACADEMIC BACKGROUND', 'OBJECTIVE', 'KEY SKILLS'
                ]
                is_section_heading = is_section_heading or stripped.upper() in common_sections
                
                if is_section_heading:
                    run = p.add_run(stripped)
                    run.bold = True
                    run.font.size = Pt(15)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.space_before = Pt(20)
                    if stripped.upper() == "PROJECTS":
                        in_projects_section = True
                    else:
                        in_projects_section = False
                else:
                    is_project_title = (
                        in_projects_section and 
                        not stripped.startswith(('•', '-', '*', '○')) and 
                        not stripped.startswith(('http', 'www', 'Tech Stack', 'Languages', 'Frameworks', 'Databases', 'Tools', 'Concepts', 'Cloud'))
                    )
                    if is_project_title:
                        run = p.add_run(stripped)
                        run.bold = True
                        run.font.size = Pt(13)
                        p.paragraph_format.space_before = Pt(10)
                        p.paragraph_format.space_after = Pt(4)
                    else:
                        p.add_run(stripped)
                        p.paragraph_format.line_spacing = 1.2
                        p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)

    doc.save(buffer)
    buffer.seek(0)
    return buffer
